import streamlit as st
import requests
from PIL import Image
from io import BytesIO
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
import zipfile

st.set_page_config(page_title="SPOT - Automação TRN", layout="wide")

# --- Funções auxiliares ---

def inserir_logo():
    logo_path = "./logo_spot.png"
    logo_img = Image.open(logo_path)
    st.image(logo_img, width=250)


def extrair_links_e_ids(file):
    wb = load_workbook(file, data_only=True)
    ws = wb.active
    headers = {cell.value: idx for idx, cell in enumerate(ws[1])}

    required = ["Link do Anexo", "ID da Despesa", "ID do Relatório"]
    if any(col not in headers for col in required):
        raise ValueError("A planilha deve conter as colunas 'Link do Anexo', 'ID da Despesa' e 'ID do Relatório'.")

    col_link = headers["Link do Anexo"]
    col_id_despesa = headers["ID da Despesa"]
    col_id_relatorio = headers["ID do Relatório"]
    dados = []

    for row in ws.iter_rows(min_row=2):
        linha_excel = row[0].row
        id_despesa = row[col_id_despesa].value
        id_relatorio = row[col_id_relatorio].value
        cell_link = row[col_link]
        url = cell_link.hyperlink.target if cell_link.hyperlink else None
        dados.append((linha_excel, id_despesa, id_relatorio, url))

    return dados


def pdf_para_imagens(pdf_bytes):
    imagens = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            img = Image.open(BytesIO(pix.tobytes("png")))
            imagens.append(img)
    return imagens


def ajustar_altura_doc_paragrafo(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepLines'))
    pPr.append(OxmlElement('w:keepNext'))


def inserir_imagem_redimensionada(paragraph, img, largura_max=5.5, altura_max=7):
    img_io = BytesIO()
    img.save(img_io, format='PNG')
    img_io.seek(0)
    largura, altura = img.size
    escala = min((largura_max * 96) / largura, (altura_max * 96) / altura) * 1.1
    nova_largura = largura * escala / 96
    run = paragraph.add_run()
    run.add_picture(img_io, width=Inches(nova_largura))


def aplicar_fonte_arial(run):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)

# --- Interface ---

modo = st.sidebar.radio("ESCOLHA A FUNCIONALIDADE:", ["Montar evidências", "Download de imagens"])

inserir_logo()
st.title("Automação TRN - Evidências VExpenses")

uploaded_file = st.file_uploader("📂 Envie a planilha (.xlsx)", type=["xlsx"])
manual_uploads = {}

if uploaded_file:
    try:
        info_links = extrair_links_e_ids(uploaded_file)

        # ---- Montar evidências ----
        if modo == "Montar evidências":
            if not info_links:
                st.error("❌ Nenhum link encontrado na planilha.")
            else:
                st.success(f"✅ {len(info_links)} registros encontrados.")

                # Verifica se há anexos ausentes
                missing = [(linha, id_despesa, id_relatorio) for linha, id_despesa, id_relatorio, url in info_links if not url]
                if missing:
                    st.header("📸 Upload manual para anexos ausentes")
                    for linha, id_despesa, id_relatorio in missing:
                        st.warning(f"🔍 Imagem ausente: ID da Despesa {id_despesa} | ID do Relatório {id_relatorio}")
                        img = st.file_uploader(
                            f"Envie a imagem para linha {linha}",
                            type=["jpg", "png", "jpeg"],
                            key=f"upload_linha_{linha}_despesa_{id_despesa}"
                        )
                        manual_uploads[linha] = img

                # Pausa se ainda faltam uploads manuais
                imagens_pendentes = [linha for linha, _, _, url in info_links if not url and not manual_uploads.get(linha)]
                if imagens_pendentes:
                    st.info("⏳ Aguardando envio de todas as imagens manuais antes de gerar o Word.")
                    st.stop()

                # Botão para gerar Word
                if st.button("📝 Gerar Documento Word"):
                    erros = []
                    doc = Document()
                    log_area = st.empty()

                    for i, (linha, id_despesa, id_relatorio, url) in enumerate(info_links, 1):
                        try:
                            log_area.markdown(f"🔄 Processando linha **{linha}**: ID da Despesa `{id_despesa}` / ID do Relatório `{id_relatorio}`")

                            if not url:
                                img = Image.open(manual_uploads.get(linha)).convert("RGB")
                                imagens = [img]
                            else:
                                if not url.startswith("http"):
                                    url = "https://" + url
                                resp = requests.get(url, timeout=20)
                                resp.raise_for_status()
                                ct = resp.headers.get('Content-Type', '')

                                if 'pdf' in ct:
                                    imagens = pdf_para_imagens(resp.content)
                                else:
                                    img = Image.open(BytesIO(resp.content)).convert("RGB")
                                    extrema = img.getextrema()
                                    if all(e[0] == e[1] for e in extrema):
                                        raise ValueError("Imagem aparentemente em branco.")
                                    imagens = [img]

                            for img in imagens:
                                doc.add_page_break()
                                p = doc.add_paragraph()
                                ajustar_altura_doc_paragrafo(p)
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = p.add_run(f"ID da Despesa: {id_despesa} / ID do Relatório: {id_relatorio}")
                                aplicar_fonte_arial(run)

                                p_img = doc.add_paragraph()
                                p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                inserir_imagem_redimensionada(p_img, img)

                        except Exception as e:
                            erros.append((linha, id_despesa, id_relatorio, e))
                            doc.add_page_break()
                            p = doc.add_paragraph()
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run(f"ID da Despesa: {id_despesa} / ID do Relatório: {id_relatorio}")
                            aplicar_fonte_arial(run)
                            p_err = doc.add_paragraph()
                            p_err.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run_err = p_err.add_run("⚠️ Erro ao carregar imagem: " + str(e))
                            aplicar_fonte_arial(run_err)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    log_area.empty()

                    st.success("✅ Documento Word gerado com sucesso!")
                    st.download_button(
                        label="📥 Baixar Word",
                        data=buffer,
                        file_name="anexos_ordenados.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    if erros:
                        st.markdown("### ❌ Falhas detectadas")
                        for linha, id_despesa, id_relatorio, erro in erros:
                            st.write(f"Linha {linha} | Despesa: {id_despesa} | Relatório: {id_relatorio} → {erro}")

        # ---- Download de imagens ----
        elif modo == "Download de imagens":
            if st.button("📥 Baixar Imagens"):
                images_data = []
                erros = []
                status_area = st.empty()

                for i, (linha, id_despesa, id_relatorio, url) in enumerate(info_links, 1):
                    try:
                        status_area.markdown(f"🔄 Processando linha **{linha}**: ID da Despesa `{id_despesa}` / ID do Relatório `{id_relatorio}`")
                        if not url:
                            raise ValueError("URL ausente")
                        if not url.startswith("http"):
                            url = "https://" + url
                        resp = requests.get(url, timeout=20)
                        resp.raise_for_status()
                        ct = resp.headers.get('Content-Type', '')

                        if 'pdf' in ct:
                            imgs = pdf_para_imagens(resp.content)
                        else:
                            img = Image.open(BytesIO(resp.content)).convert("RGB")
                            imgs = [img]

                        for idx, img in enumerate(imgs, 1):
                            buf = BytesIO()
                            img.save(buf, format='PNG')
                            data = buf.getvalue()
                            name = f"{id_despesa}.png" if len(imgs) == 1 else f"{id_despesa}_{idx}.png"
                            images_data.append((name, data))

                    except Exception as e:
                        erros.append((linha, id_despesa, e))

                status_area.empty()

                if not images_data:
                    st.error("❌ Nenhuma imagem processada.")
                else:
                    # Agrupa em sub-ZIPs de até 200MB
                    subzips = []
                    part = 1
                    current_buf = BytesIO()
                    current_zip = zipfile.ZipFile(current_buf, 'w', zipfile.ZIP_DEFLATED)
                    current_size = 0
                    max_size = 200 * 1024 * 1024

                    for filename, data in images_data:
                        if current_size + len(data) > max_size:
                            current_zip.close()
                            subzips.append((f"imagens_part{part}.zip", current_buf.getvalue()))
                            part += 1
                            current_buf = BytesIO()
                            current_zip = zipfile.ZipFile(current_buf, 'w', zipfile.ZIP_DEFLATED)
                            current_size = 0
                        current_zip.writestr(filename, data)
                        current_size += len(data)

                    current_zip.close()
                    subzips.append((f"imagens_part{part}.zip", current_buf.getvalue()))

                    # ZIP mestre
                    master_buf = BytesIO()
                    master_zip = zipfile.ZipFile(master_buf, 'w', zipfile.ZIP_DEFLATED)
                    for name, data in subzips:
                        master_zip.writestr(name, data)
                    master_zip.close()
                    master_buf.seek(0)

                    st.success("✅ Arquivo ZIP gerado com sucesso!")
                    st.download_button(
                        label="📥 Baixar Imagens Compactadas", 
                        data=master_buf,
                        file_name="imagens_compactadas.zip",
                        mime="application/zip"
                    )

                    if erros:
                        st.warning("📌 Algumas falhas ocorreram:")
                        for linha, id_despesa, erro in erros:
                            st.write(f"Linha {linha} | Despesa {id_despesa} → {erro}")

    except Exception as e:
        st.error(f"Erro ao processar: {e}")